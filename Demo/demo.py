import streamlit as st
import os
import tempfile
import ollama
import docx
import time
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder
from langchain_text_splitters import RecursiveCharacterTextSplitter
from streamlit.runtime.uploaded_file_manager import UploadedFile
from system_prompt import system_prompt
import chromadb
from chromadb.utils.embedding_functions.ollama_embedding_function import OllamaEmbeddingFunction

USER_AVATAR = "assets/user.png"
ASSISTANT_AVATAR = "assests/model.png"

st.set_page_config(page_title="Local RAG Chat", page_icon="✨")
st.title("✨ Local RAG Chat")
st.divider()

# -------------------- Session State Defaults --------------------
st.session_state.setdefault("temperature", 0.7)
st.session_state.setdefault("top_p", 0.9)
st.session_state.setdefault("max_tokens", 512)
st.session_state.setdefault("messages", [])
st.session_state.setdefault("selected_model", "")
st.session_state.setdefault("system_prompt", system_prompt)

# -------------------- Helper Functions --------------------

@st.cache_resource
def get_reranker():
    return CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")

@st.cache_resource
def get_chroma_client():
    return chromadb.PersistentClient(path="./demo-rag-app")

def get_available_models():
    try:
        models_response = ollama.list()
        return [m.model for m in models_response.models]
    except Exception as e:
        return []

def call_llm(context, prompt, model_name):
    current_sys_prompt = st.session_state.get("system_prompt", system_prompt)

    response = ollama.chat(
        model=model_name,
        stream=True,
        messages=[
            {"role": "system", "content": current_sys_prompt},
            {"role": "user", "content": f"Context: {context}. Question: {prompt}"},
        ],
        options={
            "temperature": st.session_state.temperature,
            "top_p": st.session_state.top_p,
            "num_predict": st.session_state.max_tokens
        }
    )
    
    full_metadata = {"eval_count": 0, "prompt_eval_count": 0}
    for chunk in response:
        if "message" in chunk and "content" in chunk["message"]:
            yield chunk["message"]["content"]
        
        if chunk.get("done"):
            full_metadata["eval_count"] = chunk.get("eval_count", 0)
            full_metadata["prompt_eval_count"] = chunk.get("prompt_eval_count", 0)
            st.session_state["last_metadata"] = full_metadata

def get_vector_collection():
    ollama_ef = OllamaEmbeddingFunction(url="http://localhost:11434", model_name="nomic-embed-text:v1.5")
    client = get_chroma_client()
    return client.get_or_create_collection(
        name="rag-app",
        embedding_function=ollama_ef,
        metadata={"hnsw:space": "cosine"},
    )

def add_to_collection(all_splits, file_name):
    collection = get_vector_collection()
    documents, metadata, ids = [], [], []
    for idx, split in enumerate(all_splits):
        documents.append(split.page_content)
        metadata.append(split.metadata)
        ids.append(f"{file_name}_{idx}")
    collection.upsert(documents=documents, metadatas=metadata, ids=ids)
    st.toast(f"Knowledge Base Updated!", icon="✅")

def process_pdf(file: UploadedFile):
    with tempfile.NamedTemporaryFile("wb", suffix=".pdf", delete=False) as temp_file:
        temp_file.write(file.read())
        temp_path = temp_file.name
    try:
        loader = PyMuPDFLoader(temp_path)
        docs = loader.load()
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
    return split_documents(docs)

def read_docx(file: UploadedFile):
    try:
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX file: {e}"

def read_txt(file: UploadedFile):
    try:
        return file.getvalue().decode("utf-8")
    except Exception as e:
        return f"Error reading text file: {e}"

def split_documents(docs):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, 
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", "?", "!", " ", ""]
    )
    return text_splitter.split_documents(docs)

def re_rank_cross_encoders(documents, prompt):
    encoder_model = get_reranker()
    ranks = encoder_model.rank(prompt, documents, top_k=3)
    relevant_text = "\n\n".join([documents[r["corpus_id"]] for r in ranks])
    relevant_ids = [r["corpus_id"] for r in ranks]
    return relevant_text, relevant_ids

def query_collection(prompt, n_results=10):
    collection = get_vector_collection()
    return collection.query(query_texts=[prompt], n_results=n_results)

def clear_chat_and_database():
    st.session_state.messages = []
    client = get_chroma_client()
    try:
        client.delete_collection("rag-app")
    except:
        pass 
    get_vector_collection()
    st.success("Database cleared!")


st.markdown("""
<style>
    /* Global App Styling */
    .stApp {
        background-color: #212121;
        color: #FAFAFA;
    }

    /* Centering the main content area */
    .block-container {
        max-width: 850px;
        padding-top: 2rem;
        padding-bottom: 10rem;
    }

    /* Sidebar Styling */
    [data-testid="stSidebar"] {
        background-color: #191919;
        border-right: 1px solid #333333;
    }
    

    /* Input Container Centering */
    .stChatInputContainer {
        max-width: 950px;
        margin: 0 auto;
        background-color: #212121;
    }
    
    /* Chat Messages Polishing */
    .stChatMessage {
        background-color: #2b2b2b;
        border: 1px solid #3d3d3d;
        border-radius: 12px;
        margin-bottom: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.2);
        position: relative;
    }
    
    .stChatMessage [data-testid="stMarkdownContainer"] p {
        color: #FAFAFA;
    }

    /* Custom Stats Bar Styling */
    .stats-container {
        display: flex;
        justify-content: space-between;
        gap: 10px;
        margin: 15px 0;
        margin-right: 12px;
        padding: 12px;
        background-color: #191919;
        border-left: 4px solid #DB3434;
        border-radius: 6px;
        font-family: 'Courier New', monospace;
        font-size: 0.85rem;
        color: #FAFAFA;
    }

    /* Expander Styling */
    .streamlit-expanderHeader {
        background-color: #191919 !important;
        color: #FAFAFA !important;
        border: 1px solid #333333 !important;
    }

    /* Buttons Styling */
    .stButton>button {
        border-radius: 6px;
        border: 1px solid #3d3d3d;
        transition: all 0.2s ease;
    }
    
    .stButton>button:hover {
        border-color: #DB3434 !important;
        color: #DB3434 !important;
    }
</style>
""", unsafe_allow_html=True)

with st.sidebar:    
    st.subheader("Model Selection")
    models = get_available_models()

    current_model_idx = models.index(st.session_state.selected_model) if st.session_state.selected_model in models else 0
    st.session_state.selected_model = st.selectbox(
        "Select Ollama Model", 
        options=models, 
        index=current_model_idx,
        key="model_select"
    )
    
    if st.session_state.selected_model:
        st.success(f"Active Model: **{st.session_state.selected_model}**")
    else:
        st.error("No models found. Please check Ollama installation.")
    
    st.divider()

    st.header("📚 Knowledge Base")
    uploaded_file = st.file_uploader("Upload Document", type=["pdf", "docx", "txt", "md"])
    
    if uploaded_file:
        normalize_name = uploaded_file.name.translate(str.maketrans({"-":"_", ".":"_"}))
        file_name = uploaded_file.name

        if file_name.endswith(".pdf"):
            all_splits = process_pdf(uploaded_file)
        elif file_name.endswith(".docx"):
            text = read_docx(uploaded_file)
            if "Error" in text: st.error(text)
            else: all_splits = split_documents([Document(page_content=text, metadata={"source": file_name})])
        elif file_name.endswith((".txt", ".md")):
            text = read_txt(uploaded_file)
            if "Error" in text: st.error(text)
            else: all_splits = split_documents([Document(page_content=text, metadata={"source": file_name})])
        else:
            all_splits = []

        if all_splits and st.button("Process Document", use_container_width=True):
            add_to_collection(all_splits, normalize_name)

    if st.button("Clear Database", type="secondary", use_container_width=True):
        clear_chat_and_database()
        st.rerun()

    if st.button("Clear Chat History", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

    st.divider()

    on = st.toggle("Advanced Parameters", help="Control model creativity and response length")
    if on:
        with st.expander("🛠️ Advanced Parameters", expanded=True):
            st.session_state.temperature = st.slider("Temperature", 0.0, 1.0, st.session_state.temperature, 0.01)
            st.session_state.top_p = st.slider("Top P", 0.0, 1.0, st.session_state.top_p, 0.01)
            st.session_state.max_tokens = st.number_input("Max Tokens", 64, 8192, st.session_state.max_tokens, 64)


for message in st.session_state.messages:
    avatar = USER_AVATAR if message["role"] == "user" else ASSISTANT_AVATAR
    with st.chat_message(message["role"], avatar=avatar):
        st.markdown(message["content"])

prompt = st.chat_input("What can I help you with?")

if prompt:
    with st.chat_message("user", avatar=USER_AVATAR):
        st.markdown(prompt)
    st.session_state.messages.append({"role":"user","content":prompt})
    
    with st.status("Searching knowledge base...", expanded=False) as status:
        results = query_collection(prompt)
        
        if results and results.get("documents") and len(results["documents"][0]) > 0:
            context_list = results.get("documents")[0]
            relevant_text, relevant_ids = re_rank_cross_encoders(context_list, prompt)
            status.update(label="Information retrieved and reranked.", state="complete")
        else:
            relevant_text = "No relevant context found in the database."
            relevant_ids = []
            status.update(label="No relevant information found.", state="error")

    with st.chat_message("assistant", avatar=ASSISTANT_AVATAR):
        start_time = time.time()
        
        response_placeholder = st.empty()
        full_response = ""
        
        for chunk in call_llm(context=relevant_text, prompt=prompt, model_name=st.session_state.selected_model):
            full_response += chunk
            response_placeholder.markdown(full_response + "▌")
        
        response_placeholder.markdown(full_response)
        end_time = time.time()
        
        st.session_state.messages.append({"role": "assistant", "content": full_response})

        metadata = st.session_state.get("last_metadata", {})
        total_tokens = metadata.get("eval_count", 0)
        time_taken = (end_time - start_time)
        throughput = (total_tokens / time_taken) if time_taken > 0 else 0

        st.markdown(f"""
        <div class="stats-container">
            <span>⏱️Time_Taken: {time_taken:.2f}s</span>
            <span>📊Total_Tokens:{total_tokens} tokens</span>
            <span>🚀Thorughput:{throughput:.1f} t/s</span>
        </div>
        """, unsafe_allow_html=True)

# with st.expander("📄 Document Sources & Chunks"):
#     if results and results.get("documents"):
#         st.write(f"**Reranking IDs:** {relevant_ids}")
#         st.divider()
#         for idx, doc in enumerate(results["documents"][0][:3]):
#             st.caption(f"Chunk {idx+1}")
#             st.info(doc)
#     else:
#         st.write("None")