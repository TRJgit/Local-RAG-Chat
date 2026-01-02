import streamlit as st
import os
import tempfile
import ollama
import docx
import time
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder, SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter
from streamlit.runtime.uploaded_file_manager import UploadedFile
import chromadb
from chromadb.api.types import EmbeddingFunction, Documents, Embeddings

USER_AVATAR = "assets/user.png"
ASSISTANT_AVATAR = "assets/model.png"

def read_system_prompt():
    with open("system_prompt.txt", "r") as f:
        return f.read()

system_prompt = read_system_prompt()
st.set_page_config(page_title="Local RAG Chat")
st.title("Local RAG Chat")
st.divider()


st.session_state.setdefault("temperature", 0.7)
st.session_state.setdefault("top_p", 0.9)
st.session_state.setdefault("max_tokens", 512)
st.session_state.setdefault("messages", [])
st.session_state.setdefault("selected_model", "")
st.session_state.setdefault("system_prompt", system_prompt)


@st.cache_resource
def get_reranker():
    return CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")

@st.cache_resource
def get_embedding_model():
    return SentenceTransformer("nomic-ai/nomic-embed-text-v1", trust_remote_code=True)

@st.cache_resource
def get_chroma_client():
    return chromadb.PersistentClient(path="./demo-rag-app")


class SentenceTransformerEmbeddingFunction(EmbeddingFunction):
    """Custom embedding function using SentenceTransformer"""
    def __init__(self, model):
        self.model = model
    
    def __call__(self, input: Documents) -> Embeddings:
        embeddings = self.model.encode(input)
        return embeddings


def get_available_models():
    try:
        models_response = ollama.list()
        return [m.model for m in models_response.models]
    except Exception as e:
        return []

def call_llm(context, prompt, model_name):
    current_sys_prompt = st.session_state.get("system_prompt", system_prompt)
    
    messages = [{"role": "system", "content": current_sys_prompt}]
    
    for msg in st.session_state.messages:
        messages.append({"role": msg["role"], "content": msg["content"]})
    
    messages.append({
        "role": "user",
        "content": f"Context: {context}. Question: {prompt}"
    })
    
    response = ollama.chat(
        model=model_name,
        stream=True,
        messages=messages,
        options={
            "temperature": st.session_state.temperature,
            "top_p": st.session_state.top_p,
            "num_predict": st.session_state.max_tokens
        }
    )
    
    full_metadata = {"eval_count": 0, "prompt_eval_count": 0}
    for chunk in response:
        if "message" in chunk and "content" in chunk["message"]:
            yield chunk["message"]["content"]
        
        if chunk.get("done"):
            full_metadata["eval_count"] = chunk.get("eval_count", 0)
            full_metadata["prompt_eval_count"] = chunk.get("prompt_eval_count", 0)
            st.session_state["last_metadata"] = full_metadata

def get_vector_collection():
    embedding_function = SentenceTransformerEmbeddingFunction(get_embedding_model())
    client = get_chroma_client()
    return client.get_or_create_collection(
        name="rag-app",
        embedding_function=embedding_function,
        metadata={"hnsw:space": "cosine"},
    )

def add_to_collection(all_splits, file_name):
    collection = get_vector_collection()
    documents, metadata, ids = [], [], []
    for idx, split in enumerate(all_splits):
        documents.append(split.page_content)
        metadata.append(split.metadata)
        ids.append(f"{file_name}_{idx}")
    collection.upsert(documents=documents, metadatas=metadata, ids=ids)
    st.toast(f"Knowledge Base Updated!", icon="✅")

def process_pdf(file: UploadedFile):
    with tempfile.NamedTemporaryFile("wb", suffix=".pdf", delete=False) as temp_file:
        temp_file.write(file.read())
        temp_path = temp_file.name
    try:
        loader = PyMuPDFLoader(temp_path)
        docs = loader.load()
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
    return split_documents(docs)

def read_docx(file: UploadedFile):
    try:
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX file: {e}"

def read_txt(file: UploadedFile):
    try:
        return file.getvalue().decode("utf-8")
    except Exception as e:
        return f"Error reading text file: {e}"

def split_documents(docs):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, 
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", "?", "!", " ", ""]
    )
    return text_splitter.split_documents(docs)

def re_rank_cross_encoders(documents, prompt):
    encoder_model = get_reranker()
    ranks = encoder_model.rank(prompt, documents, top_k=3)
    relevant_text = "\n\n".join([documents[r["corpus_id"]] for r in ranks])
    relevant_ids = [r["corpus_id"] for r in ranks]
    return relevant_text, relevant_ids

def query_collection(prompt, n_results=10):
    collection = get_vector_collection()
    return collection.query(query_texts=[prompt], n_results=n_results)

def clear_chat_and_database():
    st.session_state.messages = []
    client = get_chroma_client()
    try:
        client.delete_collection("rag-app")
    except:
        pass 
    get_vector_collection()
    st.success("Database cleared!")


def app_styling(file_name):
    with open(file_name) as file:
        st.markdown(f"<style>{file.read()}</style>", unsafe_allow_html=True)


with st.sidebar:    
    st.header("Model Selection")
    models = get_available_models()
    current_model_idx = models.index(st.session_state.selected_model) if st.session_state.selected_model in models else 0
    st.session_state.selected_model = st.selectbox(
        "Select Ollama Model", 
        options=models, 
        index=current_model_idx,
        key="model_select"
    )
    
    if st.session_state.selected_model:
        st.success(f"Active Model: **{st.session_state.selected_model}**")
    else:
        st.error("No models found. Please check Ollama installation.")
    
    st.divider()

    st.header("Upload Document")
    uploaded_file = st.file_uploader("Select a file", type=["pdf", "docx", "txt", "md"])
    
    if uploaded_file:
        normalize_name = uploaded_file.name.translate(str.maketrans({"-":"_", ".":"_"}))
        file_name = uploaded_file.name
        if file_name.endswith(".pdf"):
            all_splits = process_pdf(uploaded_file)
        elif file_name.endswith(".docx"):
            text = read_docx(uploaded_file)
            if "Error" in text: st.error(text)
            else: all_splits = split_documents([Document(page_content=text, metadata={"source": file_name})])
        elif file_name.endswith((".txt", ".md")):
            text = read_txt(uploaded_file)
            if "Error" in text: st.error(text)
            else: all_splits = split_documents([Document(page_content=text, metadata={"source": file_name})])
        else:
            all_splits = []
        if all_splits and st.button("Process Document", use_container_width=True):
            add_to_collection(all_splits, normalize_name)

    if st.button("Clear Database", type="secondary", use_container_width=True):
        clear_chat_and_database()
        st.toast("Database cleared successfully!")
        time.sleep(0.5)
        st.rerun()

    if st.button("Clear Chat History", use_container_width=True):
        st.session_state.messages = []
        st.toast("Chat history cleared!")
        time.sleep(0.5)
        st.rerun()

    st.divider()

    on = st.toggle("Advanced Parameters", help="Control model creativity and response length")
    if on:
        with st.expander("Advanced Parameters", expanded=True):
            st.session_state.temperature = st.slider("Temperature", 0.0, 1.0, st.session_state.temperature, 0.01)
            st.session_state.top_p = st.slider("Top P", 0.0, 1.0, st.session_state.top_p, 0.01)
            st.session_state.max_tokens = st.number_input("Max Tokens", 64, 8192, st.session_state.max_tokens, 64)

for message in st.session_state.messages:
    avatar = USER_AVATAR if message["role"] == "user" else ASSISTANT_AVATAR
    with st.chat_message(message["role"], avatar=avatar):
        st.markdown(message["content"])

prompt = st.chat_input("What can I help you with?")

if prompt:
    with st.chat_message("user", avatar=USER_AVATAR):
        st.markdown(prompt)
    st.session_state.messages.append({"role":"user","content":prompt})
    
    with st.status("Searching knowledge base...", expanded=False) as status:
        results = query_collection(prompt)
        
        if results and results.get("documents") and len(results["documents"][0]) > 0:
            context_list = results.get("documents")[0]
            relevant_text, relevant_ids = re_rank_cross_encoders(context_list, prompt)
            status.update(label="Information retrieved and reranked.", state="complete")
        else:
            relevant_text = "No relevant context found in the database."
            relevant_ids = []
            status.update(label="No relevant information found.", state="error")
    
    with st.chat_message("assistant", avatar=ASSISTANT_AVATAR):
        # start_time = time.time()
        
        response_placeholder = st.empty()
        full_response = ""
        
        for chunk in call_llm(context=relevant_text, prompt=prompt, model_name=st.session_state.selected_model):
            full_response += chunk
            response_placeholder.markdown(full_response + "▌")
        
        response_placeholder.markdown(full_response)
        # end_time = time.time()
        
        st.session_state.messages.append({"role": "assistant", "content": full_response})

        # Performance numerics (Time taken, Tokens used, Thorughput) 
        # metadata = st.session_state.get("last_metadata", {})
        # total_tokens = metadata.get("eval_count", 0)
        # time_taken = (end_time - start_time)
        # throughput = (total_tokens / time_taken) if time_taken > 0 else 0
        # st.markdown(f"""
        # <div class="stats-container">
        #     <span>Time_Taken: {time_taken:.2f}s</span>
        #     <span>Total_Tokens: {total_tokens} tokens</span>
        #     <span>Throughput: {throughput:.1f} t/s</span>
        # </div>
        # """, unsafe_allow_html=True)